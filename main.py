"""
Cyber Audit Tool — AI + NLP Edition
Tabs: System Info | Security Checks | Vulnerabilities | Log Analysis | AI Insights | Audit Report
AI:  Isolation Forest anomaly detection  +  NLP log classifier  +  trend graphs
"""

import requests
import tkinter as tk
from tkinter import filedialog, messagebox
from docx import Document
import platform, psutil, socket, subprocess, datetime, threading, os, re, sys, sqlite3

# ── Optional AI / NLP imports ─────────────────────────────────────────────────
try:
    import numpy as np
    from sklearn.ensemble import IsolationForest
    from sklearn.feature_extraction.text import TfidfVectorizer
    from sklearn.naive_bayes import MultinomialNB
    from sklearn.pipeline import Pipeline
    import matplotlib
    matplotlib.use("TkAgg")
    from matplotlib.figure import Figure
    from matplotlib.backends.backend_tkagg import FigureCanvasTkAgg
    AI_AVAILABLE = True
except ImportError:
    AI_AVAILABLE = False

# ═══════════════════════════════════════════════════════════════════════════════
#  COLOUR PALETTE
# ═══════════════════════════════════════════════════════════════════════════════
C = {
    "navy":        "#1a3a6b",
    "navy_dark":   "#122d58",
    "navy_light":  "#1e4480",
    "white":       "#ffffff",
    "off_white":   "#f4f6fa",
    "card_bg":     "#ffffff",
    "border":      "#cdd6e8",
    "text_dark":   "#1a2a45",
    "text_mid":    "#3a5080",
    "text_light":  "#7a90b0",
    "accent":      "#2a7fff",
    "green":       "#27ae60",
    "green_bg":    "#eafaf1",
    "red":         "#e74c3c",
    "red_bg":      "#fdedec",
    "amber":       "#e67e22",
    "amber_bg":    "#fef9f0",
    "purple":      "#8e44ad",
    "purple_bg":   "#f5eef8",
    "teal":        "#1abc9c",
    "teal_bg":     "#e8faf5",
    "tab_inactive":"#cdd6e8",
    "log_bg":      "#f8f9fc",
    "log_fg":      "#3a5080",
    "divider":     "#dce4f0",
}

# ═══════════════════════════════════════════════════════════════════════════════
#  DATABASE
# ═══════════════════════════════════════════════════════════════════════════════
DB_PATH = "audit_history.db"

def init_db():
    con = sqlite3.connect(DB_PATH)
    con.execute("""
        CREATE TABLE IF NOT EXISTS snapshots (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            ts TEXT, cpu_pct REAL, ram_pct REAL,
            open_ports INTEGER, high_risk_ports INTEGER,
            fw_ok INTEGER, av_ok INTEGER, risk_score REAL
        )""")
    con.commit(); con.close()

def save_snapshot(cpu_pct, ram_pct, ports, fw_status, av_status):
    high  = len([p for p in ports if p in HIGH_RISK_PORTS])
    score = 0
    if fw_status == "DISABLED": score += 3
    elif fw_status == "PARTIAL": score += 1
    if av_status in ("DISABLED","NONE"): score += 3
    score += high * 2 + len([p for p in ports if p not in HIGH_RISK_PORTS])
    con = sqlite3.connect(DB_PATH)
    con.execute(
        "INSERT INTO snapshots VALUES (NULL,?,?,?,?,?,?,?,?)",
        (datetime.datetime.now().isoformat(), cpu_pct, ram_pct,
         len(ports), high,
         1 if fw_status=="ENABLED" else 0,
         1 if av_status=="ENABLED" else 0,
         float(score)))
    con.commit(); con.close()

def load_history(limit=50):
    con  = sqlite3.connect(DB_PATH)
    rows = con.execute(
        "SELECT ts,cpu_pct,ram_pct,open_ports,high_risk_ports,fw_ok,av_ok,risk_score "
        "FROM snapshots ORDER BY id DESC LIMIT ?", (limit,)
    ).fetchall()
    con.close()
    return list(reversed(rows))

# ═══════════════════════════════════════════════════════════════════════════════
#  PORTS / CONSTANTS
# ═══════════════════════════════════════════════════════════════════════════════
COMMON_PORTS = {
    21:"FTP", 22:"SSH", 23:"Telnet", 25:"SMTP",
    53:"DNS", 80:"HTTP", 110:"POP3", 135:"RPC",
    139:"NetBIOS", 143:"IMAP", 443:"HTTPS", 445:"SMB",
    3389:"RDP", 8080:"HTTP-Alt",
}
HIGH_RISK_PORTS   = {21, 23, 139, 445}
open_ports_global = []

PORT_RISK_INFO = {
    21:  ("HIGH",   "FTP transmits data in plaintext — credentials can be intercepted."),
    22:  ("MEDIUM", "SSH is generally safe but brute-force attacks are very common."),
    23:  ("HIGH",   "Telnet is fully unencrypted — replace with SSH immediately."),
    25:  ("MEDIUM", "SMTP open — could be used for spam relay if misconfigured."),
    53:  ("LOW",    "DNS is required for name resolution. Verify this is your DNS server."),
    80:  ("LOW",    "HTTP web server running. No sensitive data should go over plain HTTP."),
    110: ("MEDIUM", "POP3 mail retrieval — older protocol, prefer IMAPS (993)."),
    135: ("HIGH",   "RPC/DCOM — frequently exploited by worms. Close if not needed."),
    139: ("HIGH",   "NetBIOS — legacy file sharing, major attack surface on local networks."),
    143: ("MEDIUM", "IMAP mail server running. Switch to IMAPS (993) for encryption."),
    443: ("LOW",    "HTTPS web server — expected on machines running web services."),
    445: ("HIGH",   "SMB — target of EternalBlue/WannaCry. Close or firewall immediately."),
    3389:("HIGH",   "RDP — common brute-force and exploit target. Enable NLA + restrict IPs."),
    8080:("LOW",    "Alternate HTTP port. Verify which application is binding to it."),
}

# ═══════════════════════════════════════════════════════════════════════════════
#  NLP  — Log Event Classifier (TF-IDF + Naive Bayes)
# ═══════════════════════════════════════════════════════════════════════════════
_NLP_TRAINING = [
    ("An account was successfully logged on", "Login Success"),
    ("successful login user account logon", "Login Success"),
    ("logon type interactive user authenticated session", "Login Success"),
    ("A logon was attempted using explicit credentials", "Login Attempt"),
    ("failed login attempt bad password wrong", "Login Failure"),
    ("An account failed to log on bad password", "Login Failure"),
    ("Logon failure unknown username or bad password", "Login Failure"),
    ("multiple failed logon attempts same account repeated", "Brute Force"),
    ("repeated failed logon same account brute force", "Brute Force"),
    ("Special privileges assigned to new logon administrator", "Privilege Escalation"),
    ("A privileged service was called elevated rights", "Privilege Escalation"),
    ("User added to administrators group privilege escalation", "Privilege Escalation"),
    ("audit policy was changed system security settings", "Policy Change"),
    ("system audit policy modified changed configuration", "Policy Change"),
    ("network connection established outbound remote", "Network Activity"),
    ("firewall rule added modified network traffic", "Firewall Change"),
    ("windows firewall exception added rule modified", "Firewall Change"),
    ("A network share object was accessed remote", "Network Activity"),
    ("new service was installed system driver", "Service Change"),
    ("scheduled task created registered windows task", "Scheduled Task"),
    ("process created unexpected executable launched", "Suspicious Process"),
    ("powershell execution script command bypass", "Suspicious Process"),
    ("registry key modified run startup autorun", "Registry Change"),
    ("object access file opened read", "File Access"),
    ("file deleted removed object access", "File Access"),
    ("sensitive file accessed documents data", "File Access"),
    ("system shutdown restart initiated power", "System Event"),
    ("Windows is starting up boot system", "System Event"),
    ("system time was changed modified", "System Event"),
    ("audit tool started initiated scan", "System Event"),
    ("awaiting scan ready tool started", "System Event"),
    ("no windows security logs available", "System Event"),
]

_nlp_pipeline = None

def build_nlp_model():
    global _nlp_pipeline
    if not AI_AVAILABLE: return
    texts  = [t for t,_ in _NLP_TRAINING]
    labels = [l for _,l in _NLP_TRAINING]
    _nlp_pipeline = Pipeline([
        ("tfidf", TfidfVectorizer(ngram_range=(1,2), min_df=1)),
        ("clf",   MultinomialNB(alpha=0.5)),
    ])
    _nlp_pipeline.fit(texts, labels)

CATEGORY_COLORS = {
    "Login Success":        ("#27ae60", "✔"),
    "Login Attempt":        ("#e67e22", "→"),
    "Login Failure":        ("#e74c3c", "✘"),
    "Brute Force":          ("#c0392b", "🚨"),
    "Privilege Escalation": ("#8e44ad", "⬆"),
    "Policy Change":        ("#e67e22", "⚙"),
    "Firewall Change":      ("#2a7fff", "🛡"),
    "Network Activity":     ("#1abc9c", "🌐"),
    "Service Change":       ("#e67e22", "⚙"),
    "Scheduled Task":       ("#e67e22", "⏱"),
    "Suspicious Process":   ("#e74c3c", "⚠"),
    "Registry Change":      ("#e74c3c", "⚠"),
    "File Access":          ("#7a90b0", "📄"),
    "System Event":         ("#7a90b0", "ℹ"),
}

def classify_log_entry(text):
    if not AI_AVAILABLE or _nlp_pipeline is None:
        return "System Event", C["text_light"], "ℹ"
    cat      = _nlp_pipeline.predict([text])[0]
    col,icon = CATEGORY_COLORS.get(cat, (C["text_light"],"ℹ"))
    return cat, col, icon

def nlp_summarise_logs(entries):
    if not entries: return "No log entries to summarise."
    counts={}; risky=[]
    for _,_,msg in entries:
        cat,col,icon=classify_log_entry(msg)
        counts[cat]=counts.get(cat,0)+1
        if cat in ("Login Failure","Brute Force","Privilege Escalation","Suspicious Process","Registry Change"):
            risky.append(f"{icon} {cat}: {msg[:60]}")
    lines=["NLP Log Analysis Summary","─"*38]
    for cat,n in sorted(counts.items(),key=lambda x:-x[1]):
        icon=CATEGORY_COLORS.get(cat,(C["text_light"],"ℹ"))[1]
        lines.append(f"  {icon}  {cat:<25} x{n}")
    lines.append("")
    if risky:
        lines.append("Security-relevant events detected:")
        lines.extend(f"  {r}" for r in risky[:6])
    else:
        lines.append("No high-risk events found in log entries.")
    return "\n".join(lines)

# ═══════════════════════════════════════════════════════════════════════════════
#  ML ANOMALY DETECTION  (Isolation Forest)
# ═══════════════════════════════════════════════════════════════════════════════
_model         = None
_model_trained = False

def _build_feature(cpu_pct, ram_pct, ports, fw_status, av_status):
    high = len([p for p in ports if p in HIGH_RISK_PORTS])
    return [cpu_pct, ram_pct, len(ports), high,
            1 if fw_status=="ENABLED" else 0,
            1 if av_status=="ENABLED" else 0]

def train_anomaly_model():
    global _model, _model_trained
    if not AI_AVAILABLE: return False
    rows=load_history(200)
    if len(rows)<5: return False
    X=np.array([[r[1],r[2],r[3],r[4],r[5],r[6]] for r in rows])
    _model=IsolationForest(contamination=0.1,random_state=42)
    _model.fit(X); _model_trained=True; return True

def predict_anomaly(cpu_pct, ram_pct, ports, fw_status, av_status):
    if not AI_AVAILABLE or not _model_trained or _model is None:
        return False, 0.0, "Model not trained yet — need more scan history."
    feat=np.array([_build_feature(cpu_pct,ram_pct,ports,fw_status,av_status)])
    pred=_model.predict(feat)[0]; score=_model.score_samples(feat)[0]
    is_anomaly=(pred==-1); reasons=[]
    rows=load_history(200)
    if rows:
        avg_cpu=sum(r[1] for r in rows)/len(rows)
        avg_ram=sum(r[2] for r in rows)/len(rows)
        avg_port=sum(r[3] for r in rows)/len(rows)
        if cpu_pct>avg_cpu*1.5:
            reasons.append(f"CPU {cpu_pct:.1f}% is {cpu_pct-avg_cpu:.1f}% above your average ({avg_cpu:.1f}%)")
        if ram_pct>avg_ram*1.3:
            reasons.append(f"RAM {ram_pct:.1f}% is {ram_pct-avg_ram:.1f}% above your average ({avg_ram:.1f}%)")
        if len(ports)>avg_port+2:
            reasons.append(f"{len(ports)} open ports — {len(ports)-avg_port:.0f} more than your baseline")
        if fw_status!="ENABLED": reasons.append("Firewall is OFF — abnormal for this system")
        if av_status not in ("ENABLED",): reasons.append("Antivirus is OFF — abnormal for this system")
    expl="\n".join(f"  - {r}" for r in reasons) if reasons \
         else ("All metrics within normal range." if not is_anomaly else "Unusual combination detected.")
    return is_anomaly, score, expl

# ═══════════════════════════════════════════════════════════════════════════════
#  SYSTEM INFO HELPERS
# ═══════════════════════════════════════════════════════════════════════════════
def ts(): return datetime.datetime.now().strftime("%m-%d-%Y %H:%M")
def write_log(msg):
    with open("logs.txt","a") as f:
        f.write(f"{datetime.datetime.now().strftime('%H:%M:%S')} - {msg}\n")
def round_storage(sz):
    for v in [128,256,512,1024,2048]:
        if sz<=v: return v
    return sz

def get_cpu_name():
    try:
        if platform.system()=="Windows":
            return subprocess.getoutput(
                'powershell -Command "(Get-CimInstance Win32_Processor).Name"').strip()
    except: pass
    return platform.processor() or "Unknown CPU"

def get_installed_software():
    try:
        if platform.system()=="Windows":
            out=subprocess.getoutput(
                'powershell -Command "Get-ItemProperty HKLM:\\Software\\Microsoft\\Windows\\CurrentVersion\\Uninstall\\* '
                '| Select-Object DisplayName | Where-Object {$_.DisplayName} '
                '| Select-Object -First 10 | ForEach-Object { $_.DisplayName }"')
            apps=[l.strip() for l in out.splitlines() if l.strip()][:10]
            return apps if apps else ["N/A"]
    except: pass
    return ["N/A"]

def get_logged_users():
    try:
        if platform.system()=="Windows":
            out=subprocess.getoutput("query user").strip()
            users=[line.split()[0].lstrip(">") for line in out.splitlines()[1:] if line.split()]
            return ", ".join(users) if users else "N/A"
    except: pass
    return os.getlogin() if hasattr(os,"getlogin") else "N/A"

def get_uptime():
    try:
        delta=datetime.datetime.now()-datetime.datetime.fromtimestamp(psutil.boot_time())
        h,m=int(delta.total_seconds()//3600),int((delta.total_seconds()%3600)//60)
        return f"{h}h {m}m"
    except: return "N/A"

def get_disk_info():
    disks=[]
    for part in psutil.disk_partitions():
        try:
            u=psutil.disk_usage(part.mountpoint)
            disks.append({"device":part.device,"total":round_storage(round(u.total/(1024**3))),
                          "used":round(u.used/(1024**3),1),"free":round(u.free/(1024**3),1),
                          "percent":u.percent,"fs":part.fstype})
        except: continue
    return disks

def get_system_info():
    cpu=get_cpu_name(); ram=psutil.virtual_memory()
    net=psutil.net_io_counters()
    return {
        "cpu_model":  cpu,
        "cpu_short":  " ".join(cpu.split()[:4]) if cpu else "Unknown",
        "cpu_usage":  psutil.cpu_percent(interval=0.5),
        "cpu_cores":  psutil.cpu_count(logical=False),
        "cpu_threads":psutil.cpu_count(logical=True),
        "hostname":   socket.gethostname(),
        "ip":         socket.gethostbyname(socket.gethostname()),
        "os":         platform.system()+" "+platform.release(),
        "os_ver":     platform.version()[:40],
        "ram_total":  round(ram.total/(1024**3)),
        "ram_used":   round(ram.used/(1024**3),1),
        "ram_free":   round(ram.available/(1024**3),1),
        "ram_pct":    ram.percent,
        "disks":      get_disk_info(),
        "uptime":     get_uptime(),
        "users":      get_logged_users(),
        "software":   get_installed_software(),
        "arch":       platform.machine(),
        "net_sent":   round(net.bytes_sent/(1024**2),1),
        "net_recv":   round(net.bytes_recv/(1024**2),1),
    }

# ═══════════════════════════════════════════════════════════════════════════════
#  SECURITY CHECKS
# ═══════════════════════════════════════════════════════════════════════════════
def check_firewall():
    try:
        out=subprocess.getoutput("netsh advfirewall show allprofiles state")
        en=out.upper().count("ON"); dis=out.upper().count("OFF")
        if en>0 and dis==0: return "ENABLED","All profiles active"
        elif en>0:          return "PARTIAL",f"{en} of {en+dis} profiles ON"
        else:               return "DISABLED","All profiles OFF — system is exposed"
    except: return "UNKNOWN","Could not read firewall status"

def check_antivirus():
    try:
        out=subprocess.getoutput(
            'powershell -Command "Get-CimInstance -Namespace root/SecurityCenter2 '
            '-ClassName AntivirusProduct | Select displayName,productState"')
        if "displayName" not in out: return "NONE","No antivirus detected"
        name="Windows Defender" if "Defender" in out else "Third-party AV"
        match=re.search(r'\d+',out)
        enabled=bool(int(match.group())&0x10) if match else False
        return ("ENABLED" if enabled else "DISABLED"),name
    except: return "UNKNOWN","Check failed"

def check_uac():
    try:
        out=subprocess.getoutput(
            'powershell -Command "(Get-ItemProperty HKLM:\\SOFTWARE\\Microsoft\\Windows\\CurrentVersion\\Policies\\System).EnableLUA"')
        return ("ENABLED","UAC is active — apps need permission to elevate") if "1" in out \
               else ("DISABLED","UAC is OFF — apps can silently gain admin rights")
    except: return "UNKNOWN","Could not check UAC"

def check_windows_update():
    try:
        out=subprocess.getoutput(
            'powershell -Command "(New-Object -ComObject Microsoft.Update.AutoUpdate).Results.LastInstallationSuccessDate"')
        if out.strip(): return "Recent",f"Last update: {out.strip()[:19]}"
        return "Unknown","Could not determine last update date"
    except: return "Unknown","Update check skipped"

# ═══════════════════════════════════════════════════════════════════════════════
#  PORTS
# ═══════════════════════════════════════════════════════════════════════════════
def scan_port(port):
    try:
        s=socket.socket(); s.settimeout(0.4)
        if s.connect_ex(("127.0.0.1",port))==0: open_ports_global.append(port)
        s.close()
    except: pass

def scan_ports():
    global open_ports_global; open_ports_global=[]
    threads=[threading.Thread(target=scan_port,args=(p,)) for p in COMMON_PORTS]
    for t in threads: t.start()
    for t in threads: t.join()
    return sorted(open_ports_global)

# Block port ///////////
def block_high_risk_ports():
    ports = state["ports"]
    blocked = []

    for port in ports:
        if port in HIGH_RISK_PORTS:
            block_port(port)
            blocked.append(port)

    if not blocked:
        messagebox.showinfo("Info", "No high-risk ports to block.")


def block_port(port):
    try:
        cmd = f'netsh advfirewall firewall add rule name="Block Port {port}" protocol=TCP localport={port} action=block'
        subprocess.getoutput(cmd)
        messagebox.showinfo("Success", f"Port {port} blocked!")
        append_raw_log(f"Blocked port {port}", "warn")
    except Exception as e:
        messagebox.showerror("Error", str(e))

#/////////////////
        
# ═══════════════════════════════════════════════════════════════════════════════
#  LOGS
# ═══════════════════════════════════════════════════════════════════════════════
def get_recent_files():
    try:
        path=os.path.expanduser("~/Documents")
        files=sorted([os.path.join(path,f) for f in os.listdir(path)],
                     key=os.path.getmtime,reverse=True)[:8]
        return files
    except: return []

def get_event_logs():
    entries=[]
    try:
        if platform.system()=="Windows":
            out=subprocess.getoutput(
                'powershell -Command "Get-EventLog -LogName Security -Newest 15 '
                '| Select-Object TimeGenerated,EntryType,Message '
                '| ForEach-Object { $_.TimeGenerated.ToString(\'MM-dd-yyyy HH:mm\') + \'|\' '
                '+ $_.EntryType + \'|\' + $_.Message.Substring(0,[Math]::Min(80,$_.Message.Length)) }"')
            for line in out.strip().splitlines():
                if "|" in line:
                    parts=line.split("|",2)
                    if len(parts)==3: entries.append((parts[0].strip(),parts[1].strip(),parts[2].strip()))
    except: pass
    if not entries:
        entries=[(ts(),"Info","System audit tool started — awaiting scan"),
                 (ts(),"Info","No Windows Security logs available")]
    return entries

# ═══════════════════════════════════════════════════════════════════════════════
#  RISK SCORE
# # ═══════════════════════════════════════════════════════════════════════════════
def calculate_risk(fw_status,av_status,ports):
    score=0
    if fw_status=="DISABLED": score+=3
    elif fw_status=="PARTIAL": score+=1
    if av_status in ("DISABLED","NONE"): score+=3
    score+=len([p for p in ports if p in HIGH_RISK_PORTS])*2
    score+=len([p for p in ports if p not in HIGH_RISK_PORTS])
    if score<=1: return "LOW",C["green"],C["green_bg"]
    elif score<=4: return "MEDIUM",C["amber"],C["amber_bg"]
    return "HIGH",C["red"],C["red_bg"]

# ═══════════════════════════════════════════════════════════════════════════════
#  OLLAMA AI EXPLANATION
# ═══════════════════════════════════════════════════════════════════════════════
def generate_ai_explanation(info, fw, av, ports, risk):
    try:
        # Extract clean values
        fw_status = fw[0] if isinstance(fw, tuple) else str(fw)
        av_status = av[0] if isinstance(av, tuple) else str(av)
        ports_str = ", ".join(map(str, ports)) if ports else "None"

        # Build strong prompt
        prompt = f"""
You are a cybersecurity assistant.

System Analysis:
- Firewall: {fw_status}
- Antivirus: {av_status}
- Open Ports: {ports_str}
- Risk Level: {risk}

Give EXACTLY 3 bullet points:
1. What is wrong
2. Why it is dangerous
3. What should be done

Keep it short, clear, and professional.
"""

        # Send request to Ollama
        response = requests.post(
            "http://127.0.0.1:11434/api/generate",
            json={
                "model": "llama3:latest",
                "prompt": prompt.strip(),
                "stream": False,
                "options": {
                    "temperature": 0.3   # 👈 added here
                }
            },
            timeout=120
        )

        data = response.json()

        # Safety check
        if "response" not in data:
            return f"Ollama error: {data}"

        return data["response"]

    except Exception as e:
        return f"Ollama error: {str(e)}"
# ═══════════════════════════════════════════════════════════════════════════════
#  Block port ////////////////////////
# ═══════════════════════════════════════════════════════════════════════════════
def block_port(port):
    try:
        cmd = f'netsh advfirewall firewall add rule name="Block Port {port}" protocol=TCP localport={port} action=block'
        result = subprocess.getoutput(cmd)
        messagebox.showinfo("Success", f"Port {port} blocked successfully!")
        append_raw_log(f"Blocked port {port}", "warn")
    except Exception as e:
        messagebox.showerror("Error", f"Failed to block port: {e}")

# ═══════════════════════════════════════════════════════════════════════════════
#  UI SETUP
# ═══════════════════════════════════════════════════════════════════════════════
init_db()
build_nlp_model()

root=tk.Tk()
root.title("Cyber Audit Tool  🤖 AI Edition")
root.geometry("1000x680")
root.resizable(True,True)
root.configure(bg=C["navy"])

F_TITLE =("Segoe UI",11,"bold")
F_HEAD  =("Segoe UI",9,"bold")
F_BODY  =("Segoe UI",9)
F_SMALL =("Segoe UI",8)
F_MONO  =("Consolas",9)
F_MED_B =("Segoe UI",10,"bold")
F_SECT  =("Segoe UI",9,"bold")

state={"info":{},"fw":("UNKNOWN",""),"av":("UNKNOWN",""),"ports":[],"events":[]}

# ── TOP BAR ───────────────────────────────────────────────────────────────────
topbar=tk.Frame(root,bg=C["navy"],height=48); topbar.pack(fill="x",side="top")
topbar.pack_propagate(False)
dots=tk.Frame(topbar,bg=C["navy"]); dots.pack(side="left",padx=10)
for dc in ("#e74c3c","#f39c12","#2ecc71"):
    tk.Label(dots,bg=dc,width=2,relief="flat").pack(side="left",padx=2,pady=16)
tk.Label(topbar,text=" Cyber Audit Tool  🤖 AI Edition",font=F_TITLE,
         bg=C["navy"],fg=C["white"]).pack(side="left",padx=6)
btn_row=tk.Frame(topbar,bg=C["navy"]); btn_row.pack(side="right",padx=14,pady=10)
scan_btn=tk.Button(btn_row,text="  Scan System  ",font=F_HEAD,
                   bg="#2ecc71",fg="white",relief="flat",cursor="hand2",bd=0,padx=8,pady=5)
scan_btn.pack(side="left",padx=(0,8))
report_btn=tk.Button(btn_row,text="  Export Report  ",font=F_HEAD,
                     bg=C["white"],fg=C["navy"],relief="flat",cursor="hand2",bd=0,padx=8,pady=5)
report_btn.pack(side="left")

# ── TAB BAR ───────────────────────────────────────────────────────────────────
TABS=["System Info","Security Checks","Vulnerabilities","Log Analysis","AI Insights","Audit Report"]
TAB_ICONS={"System Info":"  ","Security Checks":"  ","Vulnerabilities":"  ",
            "Log Analysis":"  ","AI Insights":"  ","Audit Report":"  "}
tab_bar=tk.Frame(root,bg=C["navy_dark"],height=38); tab_bar.pack(fill="x")
tab_bar.pack_propagate(False)
tab_btns={}

def switch_tab(name):
    for n,b in tab_btns.items():
        ac=(n==name)
        abg=(C["purple"] if n=="AI Insights" else C["teal"] if n=="Audit Report" else C["accent"]) if ac else C["navy_dark"]
        b.configure(bg=abg,fg="white" if ac else C["tab_inactive"])
    for n,pf in panels.items():
        if n==name: pf.tkraise()

for i,name in enumerate(TABS):
    b=tk.Button(tab_bar,text=f"  {name}  ",font=F_HEAD,
                bg=C["accent"] if i==0 else C["navy_dark"],
                fg="white" if i==0 else C["tab_inactive"],
                relief="flat",bd=0,padx=10,pady=10,cursor="hand2",
                activebackground=C["navy_light"],activeforeground="white",
                command=lambda n=name: switch_tab(n))
    b.pack(side="left"); tab_btns[name]=b

container=tk.Frame(root,bg=C["off_white"]); container.pack(fill="both",expand=True)
panels={}
for name in TABS:
    pf=tk.Frame(container,bg=C["off_white"]); pf.place(relwidth=1,relheight=1); panels[name]=pf

# ── STATUS BAR ────────────────────────────────────────────────────────────────
sbar=tk.Frame(root,bg=C["navy_dark"],height=24); sbar.pack(fill="x",side="bottom")
sbar.pack_propagate(False)
status_lbl=tk.Label(sbar,text="Ready — click Scan System to begin",
                    font=F_SMALL,bg=C["navy_dark"],fg=C["tab_inactive"])
status_lbl.pack(side="left",padx=10)
tk.Label(sbar,text="Auto-scan: ON (every 30 min)",
         font=F_SMALL,bg=C["navy_dark"],fg=C["green"]).pack(side="right",padx=10)

def set_status(msg):
    status_lbl.configure(text=msg); root.update_idletasks()

# ── WIDGET HELPERS ────────────────────────────────────────────────────────────
def slabel(parent,text,color=None):
    tk.Label(parent,text=f"  {text}",font=F_SECT,bg=C["off_white"],
             fg=color or C["text_dark"]).pack(anchor="w",pady=(8,2))

def card(parent,padx=12,pady=4,bg=None,fill="x",expand=False,**kw):
    f=tk.Frame(parent,bg=bg or C["card_bg"],highlightbackground=C["border"],
               highlightthickness=1,**kw)
    f.pack(fill=fill,expand=expand,padx=padx,pady=pady); return f

def card_header(parent,text,bg=None,color=None):
    bg=bg or C["card_bg"]
    tk.Label(parent,text=f"  {text}",font=F_SECT,bg=bg,fg=color or C["navy"]).pack(anchor="w",pady=(8,2))
    tk.Frame(parent,bg=C["divider"],height=1).pack(fill="x",padx=8,pady=(0,4))

def kv(parent,key,val="—",vfg=None,bg=None):
    bg=bg or C["card_bg"]
    row=tk.Frame(parent,bg=bg); row.pack(fill="x",padx=10,pady=2)
    tk.Label(row,text=key,font=F_BODY,bg=bg,fg=C["text_light"],width=22,anchor="w").pack(side="left")
    lbl=tk.Label(row,text=val,font=F_BODY,bg=bg,fg=vfg or C["text_dark"],anchor="w",wraplength=300)
    lbl.pack(side="left",fill="x",expand=True); return lbl

def stat_card(parent,icon,label,val,ibg):
    f=tk.Frame(parent,bg=C["card_bg"],highlightbackground=C["border"],highlightthickness=1)
    f.pack(side="left",fill="both",expand=True,padx=(0,6))
    inner=tk.Frame(f,bg=C["card_bg"]); inner.pack(padx=10,pady=10,fill="x")
    tk.Label(inner,text=icon,font=("Segoe UI",16),bg=ibg,fg=C["navy"],width=3).pack(side="left",padx=(0,8))
    right=tk.Frame(inner,bg=C["card_bg"]); right.pack(side="left",fill="x",expand=True)
    tk.Label(right,text=label,font=F_SMALL,bg=C["card_bg"],fg=C["text_light"]).pack(anchor="w")
    vl=tk.Label(right,text=val,font=F_MED_B,bg=C["card_bg"],fg=C["text_dark"]); vl.pack(anchor="w")
    return vl

def stext(parent,height,font=None,wrap="word",bg=None):
    bg=bg or C["log_bg"]
    fr=tk.Frame(parent,bg=C["card_bg"]); fr.pack(fill="both",expand=True,padx=6,pady=(0,6))
    sb=tk.Scrollbar(fr); sb.pack(side="right",fill="y")
    t=tk.Text(fr,font=font or F_MONO,bg=bg,fg=C["log_fg"],height=height,
              relief="flat",bd=0,state="disabled",wrap=wrap,yscrollcommand=sb.set)
    t.pack(fill="both",expand=True); sb.configure(command=t.yview); return t

# ═══════════════════════════════════════════════════════════════════════════════
#  TAB 1 — SYSTEM INFO
# ═══════════════════════════════════════════════════════════════════════════════
p1=panels["System Info"]
row1=tk.Frame(p1,bg=C["off_white"]); row1.pack(fill="x",padx=12,pady=(10,4))
SC={}
SC["os"]  =stat_card(row1,"🪟","Operating System","Windows ..","#dbeafe")
SC["cpu"] =stat_card(row1,"⚙","CPU Usage","—%","#dcfce7")
SC["ram"] =stat_card(row1,"📊","RAM Usage","— GB","#fef9c3")
SC["net"] =stat_card(row1,"🌐","Network  Sent/Recv","—","#e8eaf6")

# Risk card
risk_card=tk.Frame(row1,bg=C["red_bg"],highlightbackground=C["red"],highlightthickness=2)
risk_card.pack(side="left",fill="both",expand=True)
risk_inner=tk.Frame(risk_card,bg=C["red_bg"]); risk_inner.pack(padx=12,pady=10)
tk.Label(risk_inner,text="⚠",font=("Segoe UI",18),bg=C["red_bg"],fg=C["red"]).pack(side="left",padx=(0,8))
risk_right=tk.Frame(risk_inner,bg=C["red_bg"]); risk_right.pack(side="left")
tk.Label(risk_right,text="Risk Level",font=F_SMALL,bg=C["red_bg"],fg=C["red"]).pack(anchor="w")
risk_val_lbl=tk.Label(risk_right,text="—",font=("Segoe UI",16,"bold"),bg=C["red_bg"],fg=C["red"])
risk_val_lbl.pack(anchor="w")
risk_sub_lbl=tk.Label(risk_right,text="Run a scan",font=F_SMALL,bg=C["red_bg"],fg=C["red"])
risk_sub_lbl.pack(anchor="w")

twin=tk.Frame(p1,bg=C["off_white"]); twin.pack(fill="x",padx=12,pady=4)
lc=tk.Frame(twin,bg=C["card_bg"],highlightbackground=C["border"],highlightthickness=1)
rc=tk.Frame(twin,bg=C["card_bg"],highlightbackground=C["border"],highlightthickness=1)
lc.pack(side="left",fill="both",expand=True,padx=(0,6))
rc.pack(side="left",fill="both",expand=True)
card_header(lc,"System Details")
card_header(rc,"Hardware")
SYS={}
for k in ["Hostname","IP Address","Logged Users","OS Version","Uptime"]:
    SYS[k]=kv(lc,k+":","—")
HW={}
for k in ["CPU Model","Cores / Threads","RAM Total","RAM Free","Architecture"]:
    HW[k]=kv(rc,k+":","—")

twin2=tk.Frame(p1,bg=C["off_white"]); twin2.pack(fill="x",padx=12,pady=4)
dc=tk.Frame(twin2,bg=C["card_bg"],highlightbackground=C["border"],highlightthickness=1)
sw_c=tk.Frame(twin2,bg=C["card_bg"],highlightbackground=C["border"],highlightthickness=1)
dc.pack(side="left",fill="both",expand=True,padx=(0,6))
sw_c.pack(side="left",fill="both",expand=True)
card_header(dc,"Disk Usage")
disk_inner=tk.Frame(dc,bg=C["card_bg"]); disk_inner.pack(fill="x",padx=10,pady=4)
card_header(sw_c,"Installed Software (top 10)")
sw_text=stext(sw_c,5,font=F_BODY,wrap="word")

slabel(p1,"Activity Log")
act_c=card(p1,pady=(0,6))
act_log=stext(act_c,4)

def append_syslog(msg):
    act_log.configure(state="normal")
    act_log.insert("end",msg+"\n"); act_log.see("end")
    act_log.configure(state="disabled")

# ═══════════════════════════════════════════════════════════════════════════════
#  TAB 2 — SECURITY CHECKS
# ═══════════════════════════════════════════════════════════════════════════════
p2=panels["Security Checks"]
score_top=tk.Frame(p2,bg=C["card_bg"],highlightbackground=C["border"],highlightthickness=1)
score_top.pack(fill="x",padx=12,pady=(10,6))
card_header(score_top,"Overall Security Score")
score_row=tk.Frame(score_top,bg=C["card_bg"]); score_row.pack(fill="x",padx=10,pady=(0,4))
score_pct_lbl=tk.Label(score_row,text="—%",font=("Segoe UI",30,"bold"),bg=C["card_bg"],fg=C["green"])
score_pct_lbl.pack(side="left",padx=(0,14))
score_right=tk.Frame(score_row,bg=C["card_bg"]); score_right.pack(side="left",fill="x",expand=True)
score_bar_canvas=tk.Canvas(score_right,height=14,bg=C["card_bg"],highlightthickness=0)
score_bar_canvas.pack(fill="x",pady=(6,2))
score_detail_lbl=tk.Label(score_right,text="Run a scan to calculate score",
                           font=F_SMALL,bg=C["card_bg"],fg=C["text_light"])
score_detail_lbl.pack(anchor="w")

chk_twin=tk.Frame(p2,bg=C["off_white"]); chk_twin.pack(fill="x",padx=12,pady=4)

def make_chk_card(parent,title):
    f=tk.Frame(parent,bg=C["card_bg"],highlightbackground=C["border"],highlightthickness=1)
    f.pack(side="left",fill="both",expand=True,padx=(0,6))
    card_header(f,title)
    st=tk.Label(f,text="Not checked",font=("Segoe UI",10,"bold"),bg=C["card_bg"],fg=C["text_light"])
    st.pack(anchor="w",padx=12,pady=(2,0))
    dt=tk.Label(f,text="—",font=F_SMALL,bg=C["card_bg"],fg=C["text_light"],wraplength=180,justify="left")
    dt.pack(anchor="w",padx=12,pady=(0,8))
    return st,dt

fw_st,fw_dt   = make_chk_card(chk_twin,"Firewall")
av_st,av_dt   = make_chk_card(chk_twin,"Antivirus")
uac_st,uac_dt = make_chk_card(chk_twin,"UAC")
upd_st,upd_dt = make_chk_card(chk_twin,"Windows Update")

slabel(p2,"Security Recommendations")
rec2_c=card(p2,pady=(0,8),fill="both",expand=True)
rec2_text=stext(rec2_c,12,font=F_BODY,wrap="word",bg=C["card_bg"])
for tag,col in [("H",C["red"]),("M",C["amber"]),("OK",C["green"]),("I",C["accent"])]:
    rec2_text.tag_configure(tag,foreground=col)

# ═══════════════════════════════════════════════════════════════════════════════
#  TAB 3 — VULNERABILITIES
# ═══════════════════════════════════════════════════════════════════════════════
p3=panels["Vulnerabilities"]
vuln_sum=tk.Frame(p3,bg=C["off_white"]); vuln_sum.pack(fill="x",padx=12,pady=(10,4))
VS={}
VS["total"]  =stat_card(vuln_sum,"🔍","Open Ports","0","#dbeafe")
VS["high"]   =stat_card(vuln_sum,"🚨","High Risk","0","#fdedec")
VS["medium"] =stat_card(vuln_sum,"⚠","Medium Risk","0","#fef9f0")
VS["low"]    =stat_card(vuln_sum,"ℹ","Low Risk","0","#eafaf1")

vuln_twin=tk.Frame(p3,bg=C["off_white"]); vuln_twin.pack(fill="both",expand=True,padx=12,pady=4)
port_c=tk.Frame(vuln_twin,bg=C["card_bg"],highlightbackground=C["border"],highlightthickness=1)
detail_c=tk.Frame(vuln_twin,bg=C["card_bg"],highlightbackground=C["border"],highlightthickness=1)
port_c.pack(side="left",fill="both",expand=True,padx=(0,6))
detail_c.pack(side="left",fill="both",expand=True)
card_header(port_c,"Port Scan Results (localhost)")
port_text=stext(port_c,14)
for tag,col in [("hi",C["red"]),("me",C["amber"]),("lo",C["green"]),("cl",C["text_light"])]:
    port_text.tag_configure(tag,foreground=col)
card_header(detail_c,"Vulnerability Details")
vd_text=stext(detail_c,14,font=F_BODY,wrap="word")
for tag,col in [("HIGH",C["red"]),("MEDIUM",C["amber"]),("LOW",C["green"]),
                ("h2",C["text_dark"]),("port",C["accent"])]:
    vd_text.tag_configure(tag,foreground=col)
vd_text.tag_configure("h2",font=F_HEAD)

#block port ////

block_btn = tk.Button(p3, text="Block High-Risk Ports",
                      font=("Segoe UI", 9, "bold"), bg=C["red"], fg=C["white"],
                      command=lambda: block_high_risk_ports(),
                      cursor="hand2")
block_btn.pack(pady=5)


# ═══════════════════════════════════════════════════════════════════════════════
#  TAB 4 — LOG ANALYSIS  (NLP-powered)
# ═══════════════════════════════════════════════════════════════════════════════
p4=panels["Log Analysis"]
nlp_c=tk.Frame(p4,bg=C["purple_bg"],highlightbackground=C["purple"],highlightthickness=1)
nlp_c.pack(fill="x",padx=12,pady=(10,4))
card_header(nlp_c,"NLP Log Classifier  (TF-IDF + Naive Bayes)",bg=C["purple_bg"],color=C["purple"])
nlp_sum_text=tk.Text(nlp_c,font=F_MONO,bg=C["purple_bg"],fg=C["text_dark"],
                      height=7,relief="flat",bd=0,state="disabled",wrap="word")
nlp_sum_text.pack(fill="x",padx=8,pady=(0,8))

log_twin=tk.Frame(p4,bg=C["off_white"]); log_twin.pack(fill="both",expand=True,padx=12,pady=4)
ev_c=tk.Frame(log_twin,bg=C["card_bg"],highlightbackground=C["border"],highlightthickness=1)
fi_c=tk.Frame(log_twin,bg=C["card_bg"],highlightbackground=C["border"],highlightthickness=1)
ev_c.pack(side="left",fill="both",expand=True,padx=(0,6))
fi_c.pack(side="left",fill="both",expand=True)
card_header(ev_c,"Classified Events (NLP colour-coded)")
event_text=stext(ev_c,12)
for cat,(col,icon) in CATEGORY_COLORS.items():
    event_text.tag_configure(cat,foreground=col)
event_text.tag_configure("ts",foreground=C["accent"])
event_text.tag_configure("dim",foreground=C["text_light"])
card_header(fi_c,"Recent Modified Files")
files_text=stext(fi_c,12,font=F_BODY,wrap="none")
files_text.tag_configure("new",foreground=C["accent"])
files_text.tag_configure("old",foreground=C["text_light"])

slabel(p4,"Raw Log")
raw_c=card(p4,pady=(0,8))
raw_text=stext(raw_c,4)
for tag,col in [("ts",C["accent"]),("ok",C["green"]),("warn",C["amber"]),("err",C["red"])]:
    raw_text.tag_configure(tag,foreground=col)

def append_raw_log(msg,level="ok"):
    raw_text.configure(state="normal")
    t=datetime.datetime.now().strftime("%H:%M:%S")
    raw_text.insert("end",f"[{t}]","ts"); raw_text.insert("end",f" {msg}\n",level)
    raw_text.see("end"); raw_text.configure(state="disabled")

# ═══════════════════════════════════════════════════════════════════════════════
#  TAB 5 — AI INSIGHTS
# ═══════════════════════════════════════════════════════════════════════════════
p_ai=panels["AI Insights"]

if not AI_AVAILABLE:
    mf=tk.Frame(p_ai,bg=C["off_white"]); mf.pack(fill="both",expand=True,padx=24,pady=24)
    tk.Label(mf,text="AI Libraries Not Installed",
             font=("Segoe UI",14,"bold"),bg=C["off_white"],fg=C["amber"]).pack(pady=(20,10))
    tk.Label(mf,text="Run:  pip install scikit-learn matplotlib numpy\nThen restart the app.",
             font=F_BODY,bg=C["off_white"],fg=C["text_dark"],justify="left").pack(anchor="w",padx=20)
else:
    ai_top=tk.Frame(p_ai,bg=C["off_white"]); ai_top.pack(fill="x",padx=12,pady=(10,4))
    anom_c=tk.Frame(ai_top,bg=C["card_bg"],highlightbackground=C["border"],highlightthickness=1)
    anom_c.pack(side="left",fill="both",expand=True,padx=(0,6))
    card_header(anom_c,"Isolation Forest — ML Anomaly Detection")
    anom_status=tk.Label(anom_c,text="Run a scan first",font=("Segoe UI",11,"bold"),
                          bg=C["card_bg"],fg=C["text_light"])
    anom_status.pack(anchor="w",padx=12,pady=(4,2))
    anom_detail=stext(anom_c,5,font=F_BODY,wrap="word",bg=C["card_bg"])

    stat_c=tk.Frame(ai_top,bg=C["card_bg"],highlightbackground=C["border"],highlightthickness=1)
    stat_c.pack(side="left",fill="both",expand=True)
    card_header(stat_c,"Session Statistics")
    AI_STATS={}
    for k in ["Total Scans","Avg CPU %","Avg RAM %","Anomalies Detected","NLP Categories","Model Status"]:
        AI_STATS[k]=kv(stat_c,k+":","—")

    # NLP category badge strip
    slabel(p_ai,"NLP Event Category Counts",color=C["purple"])
    badge_f=tk.Frame(p_ai,bg=C["off_white"]); badge_f.pack(fill="x",padx=12,pady=(0,4))
    nlp_badges={}
    for i,(cat,(col,icon)) in enumerate(list(CATEGORY_COLORS.items())[:8]):
        f=tk.Frame(badge_f,bg=col); f.pack(side="left",padx=(0,4),pady=2)
        tk.Label(f,text=f" {icon} {cat} ",font=F_SMALL,bg=col,fg="white").pack(side="left")
        cnt=tk.Label(f,text="0",font=("Segoe UI",8,"bold"),bg="white",fg=col,padx=4)
        cnt.pack(side="left"); nlp_badges[cat]=cnt

    slabel(p_ai,"System Trends Over Time",color=C["navy"])
    graph_frame=tk.Frame(p_ai,bg=C["off_white"]); graph_frame.pack(fill="both",expand=True,padx=12,pady=(0,4))
    fig=Figure(figsize=(9,3),dpi=88,facecolor=C["off_white"])
    fig.subplots_adjust(left=0.06,right=0.98,top=0.88,bottom=0.22,wspace=0.34)
    ax_cpu=fig.add_subplot(1,3,1); ax_ram=fig.add_subplot(1,3,2); ax_risk=fig.add_subplot(1,3,3)
    for ax,title in [(ax_cpu,"CPU %"),(ax_ram,"RAM %"),(ax_risk,"Risk Score")]:
        ax.set_title(title,fontsize=8,color=C["text_dark"])
        ax.set_facecolor(C["card_bg"])
        ax.tick_params(labelsize=6,colors=C["text_light"])
        for sp in ["bottom","left"]: ax.spines[sp].set_color(C["border"])
        for sp in ["top","right"]:   ax.spines[sp].set_visible(False)
    canvas_fig=FigureCanvasTkAgg(fig,master=graph_frame)
    canvas_fig.draw(); canvas_fig.get_tk_widget().pack(fill="both",expand=True)
    refresh_ai_btn=tk.Button(p_ai,text="  Refresh AI Analysis  ",font=F_HEAD,
                              bg=C["purple"],fg="white",relief="flat",cursor="hand2",bd=0,padx=8,pady=4)
    refresh_ai_btn.pack(pady=(2,8))

# ═══════════════════════════════════════════════════════════════════════════════
#  TAB 6 — AUDIT REPORT
# ═══════════════════════════════════════════════════════════════════════════════
p5=panels["Audit Report"]
rep_top=tk.Frame(p5,bg=C["off_white"]); rep_top.pack(fill="x",padx=12,pady=(10,4))
risk_sum_c=tk.Frame(rep_top,bg=C["card_bg"],highlightbackground=C["border"],highlightthickness=1)
risk_sum_c.pack(side="left",fill="both",padx=(0,6),ipadx=10,ipady=6)
rep_lbl=tk.Label(risk_sum_c,text="—",font=("Segoe UI",38,"bold"),bg=C["card_bg"],fg=C["red"])
rep_lbl.pack(pady=(12,2),padx=20)
rep_badge=tk.Label(risk_sum_c,text="Run scan first",font=F_SMALL,bg=C["red_bg"],fg=C["red"],padx=6,pady=2)
rep_badge.pack(pady=(0,6))
tk.Frame(risk_sum_c,bg=C["divider"],height=1).pack(fill="x",padx=8,pady=4)
REP={}
for k in ["Firewall","Antivirus","UAC","Open Ports","High-Risk Ports"]:
    REP[k]=kv(risk_sum_c,k+":","—")

rec_c=tk.Frame(rep_top,bg=C["card_bg"],highlightbackground=C["border"],highlightthickness=1)
rec_c.pack(side="left",fill="both",expand=True,ipady=6)
card_header(rec_c,"Recommendations")
rec_text=tk.Text(rec_c,font=F_BODY,bg=C["card_bg"],fg=C["text_dark"],
                 height=12,relief="flat",bd=0,state="disabled",wrap="word")
rec_text.pack(fill="both",expand=True,padx=8,pady=(0,8))
for tag,col in [("H",C["red"]),("M",C["amber"]),("OK",C["green"])]:
    rec_text.tag_configure(tag,foreground=col)

rep_bot=tk.Frame(p5,bg=C["off_white"]); rep_bot.pack(fill="both",expand=True,padx=12,pady=(0,8))
full_c=tk.Frame(rep_bot,bg=C["card_bg"],highlightbackground=C["border"],highlightthickness=1)
nlp_rep_c=tk.Frame(rep_bot,bg=C["purple_bg"],highlightbackground=C["purple"],highlightthickness=1)
full_c.pack(side="left",fill="both",expand=True,padx=(0,6))
nlp_rep_c.pack(side="left",fill="both",expand=True)
card_header(full_c,"Full Report Text")
full_text=stext(full_c,10)
card_header(nlp_rep_c,"NLP Log Summary",bg=C["purple_bg"],color=C["purple"])
nlp_rep_text=stext(nlp_rep_c,10,font=F_BODY,wrap="word",bg=C["purple_bg"])

# ═══════════════════════════════════════════════════════════════════════════════
#  POPULATE FUNCTIONS
# ═══════════════════════════════════════════════════════════════════════════════

def populate_system_tab(info):
    SC["os"].configure(text=info["os"])
    SC["cpu"].configure(text=f"{info['cpu_usage']:.0f}%")
    SC["ram"].configure(text=f"{info['ram_pct']:.0f}%  ({info['ram_used']}/{info['ram_total']} GB)")
    SC["net"].configure(text=f"Sent {info['net_sent']} MB  /  Recv {info['net_recv']} MB")
    SYS["Hostname"].configure(text=info["hostname"])
    SYS["IP Address"].configure(text=info["ip"])
    SYS["Logged Users"].configure(text=info["users"])
    SYS["OS Version"].configure(text=info["os_ver"])
    SYS["Uptime"].configure(text=info["uptime"])
    HW["CPU Model"].configure(text=info["cpu_short"])
    HW["Cores / Threads"].configure(text=f"{info['cpu_cores']} cores / {info['cpu_threads']} threads")
    HW["RAM Total"].configure(text=f"{info['ram_total']} GB")
    HW["RAM Free"].configure(text=f"{info['ram_free']} GB available")
    HW["Architecture"].configure(text=info["arch"])
    # Disk bars
    for w in disk_inner.winfo_children(): w.destroy()
    for d in info["disks"][:3]:
        col=C["green"] if d["percent"]<70 else (C["amber"] if d["percent"]<90 else C["red"])
        row=tk.Frame(disk_inner,bg=C["card_bg"]); row.pack(fill="x",pady=1)
        tk.Label(row,text=d["device"],font=F_SMALL,bg=C["card_bg"],fg=C["text_dark"],
                 width=12,anchor="w").pack(side="left")
        tk.Label(row,text=f"{d['total']} GB  ({d['percent']:.0f}% used)  FS:{d['fs']}",
                 font=F_SMALL,bg=C["card_bg"],fg=C["text_light"]).pack(side="right")
        bar=tk.Canvas(disk_inner,height=7,bg=C["card_bg"],highlightthickness=0)
        bar.pack(fill="x",pady=(0,4))
        def _draw(ev,pct=d["percent"],c=col,bc=bar):
            w=bc.winfo_width() or 180; bc.delete("all")
            bc.create_rectangle(0,0,w,7,fill=C["divider"],outline="")
            bc.create_rectangle(0,0,int(w*pct/100),7,fill=c,outline="")
        bar.bind("<Configure>",_draw); _draw(None)
    # Software
    sw_text.configure(state="normal"); sw_text.delete("1.0","end")
    SUSP={"hack","crack","keygen","inject","torrent","cheat","warez"}
    for app in info["software"][:10]:
        flag=any(w in app.lower() for w in SUSP)
        sw_text.insert("end",f"  {'! ' if flag else '  '}{app}\n","warn" if flag else "ok")
    sw_text.tag_configure("warn",foreground=C["red"])
    sw_text.tag_configure("ok",foreground=C["green"])
    sw_text.configure(state="disabled")

def populate_security_tab(fw,av,uac,upd):
    fw_st2,fw_d=fw; av_st2,av_n=av; uac_s,uac_d=uac; upd_s,upd_d=upd
    def _set(lbl,det,status,ok_vals,detail):
        ok=status in ok_vals
        col=C["green"] if ok else (C["amber"] if status in ("PARTIAL","Unknown","UNKNOWN") else C["red"])
        lbl.configure(text=f"{'OK' if ok else 'X'}  {status}",fg=col)
        det.configure(text=detail,fg=C["text_mid"])
    _set(fw_st,fw_dt,fw_st2,["ENABLED"],fw_d)
    _set(av_st,av_dt,av_st2,["ENABLED"],av_n)
    _set(uac_st,uac_dt,uac_s,["ENABLED"],uac_d)
    _set(upd_st,upd_dt,upd_s,["Recent"],upd_d)
    pts={}
    pts["fw"] =30 if fw_st2=="ENABLED" else (15 if fw_st2=="PARTIAL" else 0)
    pts["av"] =35 if av_st2=="ENABLED" else 0
    pts["uac"]=20 if uac_s=="ENABLED"  else 0
    pts["upd"]=15 if upd_s=="Recent"   else 7
    total=sum(pts.values())
    bar_col=C["green"] if total>=75 else (C["amber"] if total>=45 else C["red"])
    score_pct_lbl.configure(text=f"{total}%",fg=bar_col)
    score_detail_lbl.configure(
        text=f"Firewall {pts['fw']}/30  |  Antivirus {pts['av']}/35  |  UAC {pts['uac']}/20  |  Updates {pts['upd']}/15",
        fg=C["text_light"])
    def _draw(ev=None):
        w=score_bar_canvas.winfo_width() or 300
        score_bar_canvas.delete("all")
        score_bar_canvas.create_rectangle(0,0,w,14,fill=C["divider"],outline="")
        score_bar_canvas.create_rectangle(0,0,int(w*total/100),14,fill=bar_col,outline="")
    score_bar_canvas.bind("<Configure>",_draw); _draw()
    rec2_text.configure(state="normal"); rec2_text.delete("1.0","end")
    items=[]
    if fw_st2!="ENABLED": items.append(("H","[HIGH]   Firewall is OFF — enable on all profiles immediately.\n\n"))
    if av_st2!="ENABLED": items.append(("H","[HIGH]   Antivirus is disabled — enable real-time protection now.\n\n"))
    if uac_s!="ENABLED":  items.append(("M","[MEDIUM] UAC is disabled — apps can silently gain admin rights.\n\n"))
    if upd_s!="Recent":   items.append(("M","[MEDIUM] Windows Update status unknown — patch your system.\n\n"))
    if not items:
        items.append(("OK","All security checks passed. System looks well configured.\n\n"))
        items.append(("I","Keep running regular scans to detect configuration drift.\n\n"))
    for tag,msg in items: rec2_text.insert("end",msg,tag)
    rec2_text.tag_configure("I",foreground=C["accent"])
    rec2_text.configure(state="disabled")

def populate_vuln_tab(ports):
    port_text.configure(state="normal"); port_text.delete("1.0","end")
    vd_text.configure(state="normal"); vd_text.delete("1.0","end")
    hi=me=lo=0
    for port in sorted(COMMON_PORTS.keys()):
        name=COMMON_PORTS[port]; is_open=(port in ports)
        lvl,info_txt=PORT_RISK_INFO.get(port,("LOW","No additional info."))
        if is_open:
            tag="hi" if lvl=="HIGH" else ("me" if lvl=="MEDIUM" else "lo")
            icon="(!)" if lvl=="HIGH" else (" ! " if lvl=="MEDIUM" else " o ")
            port_text.insert("end",f" {icon} Port {port:<6} {name:<12}  [{lvl}]\n",tag)
            if lvl=="HIGH":   hi+=1
            elif lvl=="MEDIUM": me+=1
            else: lo+=1
            vd_text.insert("end",f"Port {port}  —  {name}\n","h2")
            vd_text.insert("end",f"Risk: ","h2"); vd_text.insert("end",f"{lvl}\n",lvl)
            vd_text.insert("end",f"{info_txt}\n\n")
        else:
            port_text.insert("end",f"  -  Port {port:<6} {name:<12}  [closed]\n","cl")
    port_text.configure(state="disabled")
    if not [p for p in ports if p in COMMON_PORTS]:
        vd_text.insert("end","No open ports detected.\nSystem port surface appears clean.\n")
    vd_text.configure(state="disabled")
    VS["total"].configure(text=str(len(ports)),fg=C["red"] if ports else C["green"])
    VS["high"].configure(text=str(hi),fg=C["red"] if hi else C["text_light"])
    VS["medium"].configure(text=str(me),fg=C["amber"] if me else C["text_light"])
    VS["low"].configure(text=str(lo),fg=C["green"])

def populate_log_tab(entries,files):
    event_text.configure(state="normal"); event_text.delete("1.0","end")
    cat_counts={}
    for (t,kind,msg) in entries:
        cat,col,icon=classify_log_entry(msg)
        cat_counts[cat]=cat_counts.get(cat,0)+1
        event_text.insert("end",f"[{t}] ","ts")
        event_text.insert("end",f"{icon} {cat:<22} ","dim")
        event_text.insert("end",f"{msg[:65]}\n",cat)
    event_text.configure(state="disabled")
    nlp_sum_text.configure(state="normal"); nlp_sum_text.delete("1.0","end")
    nlp_sum_text.insert("end",nlp_summarise_logs(entries))
    nlp_sum_text.configure(state="disabled")
    files_text.configure(state="normal"); files_text.delete("1.0","end")
    now=datetime.datetime.now().timestamp()
    for fp in files:
        try:
            age=(now-os.path.getmtime(fp))/3600
            tag="new" if age<24 else "old"
            sz=os.path.getsize(fp)
            szs=f"{sz//1024}KB" if sz>=1024 else f"{sz}B"
            files_text.insert("end",f"  [{'NEW' if age<24 else '   '}]  {os.path.basename(fp):<38} {szs}\n",tag)
        except:
            files_text.insert("end",f"       {os.path.basename(fp)}\n","old")
    if not files:
        files_text.insert("end","  No recent files found in Documents.\n","old")
    files_text.configure(state="disabled")
    return cat_counts

def populate_ai_tab(cpu_pct,ram_pct,ports,fw_status,av_status,cat_counts=None):
    if not AI_AVAILABLE: return
    trained=train_anomaly_model()
    is_anom,score,expl=predict_anomaly(cpu_pct,ram_pct,ports,fw_status,av_status)
    if not trained:
        anom_status.configure(text="Collecting baseline — need 5+ scans",fg=C["amber"],bg=C["card_bg"])
        anom_detail.configure(state="normal"); anom_detail.delete("1.0","end")
        anom_detail.insert("end",
            f"Isolation Forest trains on your scan history.\n"
            f"Snapshots so far: {len(load_history())}  (minimum: 5)\n\n"
            "Auto-scan fires every 30 min to build history faster.\n"
            "Or click Scan System multiple times.")
        anom_detail.configure(state="disabled")
        AI_STATS["Model Status"].configure(text="Collecting data",fg=C["amber"])
    else:
        anom_status.configure(
            text="ANOMALY DETECTED — Unusual behaviour!" if is_anom else "System behaviour is NORMAL",
            fg=C["red"] if is_anom else C["green"],bg=C["card_bg"])
        anom_detail.configure(state="normal"); anom_detail.delete("1.0","end")
        anom_detail.insert("end",
            f"Isolation Forest score: {score:.4f}  ({'ANOMALOUS' if is_anom else 'NORMAL'})\n\n"
            f"Analysis:\n{expl}\n\n"
            "Trained on historical snapshots using unsupervised ML.\nNo manual rules required.")
        anom_detail.configure(state="disabled")
        AI_STATS["Model Status"].configure(text="Trained & Active",fg=C["green"])
    rows=load_history(200)
    if rows:
        n=len(rows)
        avg_cpu=sum(r[1] for r in rows)/n; avg_ram=sum(r[2] for r in rows)/n
        n_anom=sum(1 for r in rows if _model_trained and _model is not None and
                   _model.predict(np.array([[r[1],r[2],r[3],r[4],r[5],r[6]]]))[0]==-1) if _model_trained else 0
        AI_STATS["Total Scans"].configure(text=str(n))
        AI_STATS["Avg CPU %"].configure(text=f"{avg_cpu:.1f}%")
        AI_STATS["Avg RAM %"].configure(text=f"{avg_ram:.1f}%")
        AI_STATS["Anomalies Detected"].configure(text=str(n_anom),fg=C["red"] if n_anom>0 else C["green"])
    if cat_counts:
        AI_STATS["NLP Categories"].configure(text=str(len(cat_counts)))
        for cat,lbl in nlp_badges.items():
            lbl.configure(text=str(cat_counts.get(cat,0)))
    rows=load_history(30)
    if len(rows)>=2:
        xs=list(range(len(rows))); labels=[r[0][11:16] for r in rows]
        for ax,(ys,col,title) in [(ax_cpu,([r[1] for r in rows],"#2a7fff","CPU %")),
                                   (ax_ram,([r[2] for r in rows],"#27ae60","RAM %")),
                                   (ax_risk,([r[7] for r in rows],"#e74c3c","Risk Score"))]:
            ax.clear(); ax.set_facecolor(C["card_bg"])
            ax.set_title(title,fontsize=8,color=C["text_dark"])
            ax.tick_params(labelsize=6,colors=C["text_light"])
            for sp in ["bottom","left"]: ax.spines[sp].set_color(C["border"])
            for sp in ["top","right"]:   ax.spines[sp].set_visible(False)
            ax.plot(xs,ys,color=col,linewidth=1.8,marker="o",markersize=3)
            ax.fill_between(xs,ys,alpha=0.13,color=col)
            step=max(1,len(xs)//5)
            ax.set_xticks(xs[::step]); ax.set_xticklabels(labels[::step],rotation=30,ha="right",fontsize=6)
        canvas_fig.draw()

def populate_report_tab(fw,av,ports,cat_counts=None):
    fw_status,_=fw; av_status,_=av
    risk_label,risk_col,risk_bg=calculate_risk(fw_status,av_status,ports)
    hi_ports=[p for p in ports if p in HIGH_RISK_PORTS]
    uac_s,_=check_uac()
    rep_lbl.configure(text=risk_label,fg=risk_col,bg=risk_bg)
    risk_sum_c.configure(bg=risk_bg)
    rep_badge.configure(
        text={"LOW":"SYSTEM SECURE","MEDIUM":"REVIEW NEEDED","HIGH":"ACTION REQUIRED"}.get(risk_label,""),
        bg=risk_bg,fg=risk_col)
    REP["Firewall"].configure(text=fw_status,fg=C["green"] if fw_status=="ENABLED" else C["red"])
    REP["Antivirus"].configure(text=av_status,fg=C["green"] if av_status=="ENABLED" else C["red"])
    REP["UAC"].configure(text=uac_s,fg=C["green"] if uac_s=="ENABLED" else C["amber"])
    REP["Open Ports"].configure(text=str(len(ports)),fg=C["red"] if ports else C["green"])
    REP["High-Risk Ports"].configure(text=str(len(hi_ports)),fg=C["red"] if hi_ports else C["green"])
    rec_text.configure(state="normal"); rec_text.delete("1.0","end")
    recs=[]
    if fw_status!="ENABLED": recs.append(("H","[HIGH]   Enable Windows Firewall on ALL profiles.\n\n"))
    if av_status!="ENABLED": recs.append(("H","[HIGH]   Enable Antivirus / real-time protection.\n\n"))
    for p in hi_ports:
        recs.append(("H",f"[HIGH]   Close Port {p} ({COMMON_PORTS.get(p,'')}) — {PORT_RISK_INFO.get(p,('',''))[1]}\n\n"))
    for p in [x for x in ports if x not in HIGH_RISK_PORTS]:
        recs.append(("M",f"[MEDIUM] Verify Port {p} ({COMMON_PORTS.get(p,'')}) is intentional.\n\n"))
    if uac_s!="ENABLED": recs.append(("M","[MEDIUM] Enable UAC to prevent silent privilege escalation.\n\n"))
    if not recs: recs.append(("OK","All checks passed. No critical issues found. Keep monitoring.\n\n"))
    for tag,msg in recs: rec_text.insert("end",msg,tag)
    rec_text.configure(state="disabled")
    info=state["info"]
    full_text.configure(state="normal"); full_text.delete("1.0","end")
    full_text.insert("end","\n".join([
        "="*48,
        "  CYBER AUDIT REPORT ",
        "="*48,
        f"Generated : {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}",
        f"Hostname  : {info.get('hostname','?')}  |  IP: {info.get('ip','?')}",
        f"OS        : {info.get('os','?')}",
        f"Uptime    : {info.get('uptime','?')}",
        "-"*48,
        f"Firewall  : {fw_status}",
        f"Antivirus : {av_status}",
        f"UAC       : {uac_s}",
        f"Open Ports: {', '.join(str(p) for p in ports) if ports else 'None'}",
        "-"*48,
        f"Risk Level: {risk_label}",
        f"ML Model  : {'Active (Isolation Forest)' if AI_AVAILABLE and _model_trained else 'Collecting data'}",
        f"NLP Model : {'Active (TF-IDF + Naive Bayes)' if AI_AVAILABLE else 'Not available'}",
    ]))
    full_text.configure(state="disabled")
    nlp_rep_text.configure(state="normal"); nlp_rep_text.delete("1.0","end")
    nlp_rep_text.insert("end",
        nlp_summarise_logs(state.get("events",[])) if state.get("events")
        else "Run a scan to generate NLP log summary.")
    nlp_rep_text.configure(state="disabled")
    risk_val_lbl.configure(text=risk_label,fg=risk_col)
    risk_card.configure(bg=risk_bg,highlightbackground=risk_col)
    risk_inner.configure(bg=risk_bg); risk_right.configure(bg=risk_bg)
    risk_sub_lbl.configure(
        text={"LOW":"System Secure","MEDIUM":"Review Needed","HIGH":"Action Required"}.get(risk_label,""),
        fg=risk_col,bg=risk_bg)
    for ch in [*risk_inner.winfo_children(),*risk_right.winfo_children()]:
        try: ch.configure(bg=risk_bg)
        except: pass

        
# #### NEW STYLE REPORT WORD:

# def populate_report_tab(fw,av,ports,cat_counts=None):
#     fw_status,_ = fw
#     av_status,_ = av
#     risk_label,risk_col,risk_bg = calculate_risk(fw_status,av_status,ports)

#     hi_ports = [p for p in ports if p in HIGH_RISK_PORTS]
#     uac_s,_  = check_uac()
#     info     = state["info"]

#     # ── UI CARDS (no change) ──
#     rep_lbl.configure(text=risk_label,fg=risk_col,bg=risk_bg)
#     risk_sum_c.configure(bg=risk_bg)

#     rep_badge.configure(
#         text={"LOW":"SYSTEM SECURE","MEDIUM":"REVIEW NEEDED","HIGH":"ACTION REQUIRED"}.get(risk_label,""),
#         bg=risk_bg,fg=risk_col
#     )

#     REP["Firewall"].configure(text=fw_status,fg=C["green"] if fw_status=="ENABLED" else C["red"])
#     REP["Antivirus"].configure(text=av_status,fg=C["green"] if av_status=="ENABLED" else C["red"])
#     REP["UAC"].configure(text=uac_s,fg=C["green"] if uac_s=="ENABLED" else C["amber"])
#     REP["Open Ports"].configure(text=str(len(ports)),fg=C["red"] if ports else C["green"])
#     REP["High-Risk Ports"].configure(text=str(len(hi_ports)),fg=C["red"] if hi_ports else C["green"])

#     # ── PROFESSIONAL FORM STYLE REPORT ──
#     full_text.configure(state="normal")
#     full_text.delete("1.0","end")

#     report = f"""
# ══════════════════════════════════════════════
#         CYBER SECURITY INCIDENT REPORT
# ══════════════════════════════════════════════

# Reported On : {datetime.datetime.now().strftime('%d-%m-%Y %H:%M')}
# Hostname    : {info.get('hostname','?')}
# IP Address  : {info.get('ip','?')}

# ──────────────────────────────────────────────
# SYSTEM DETAILS
# ──────────────────────────────────────────────
# Operating System : {info.get('os','?')}
# CPU              : {info.get('cpu_model','?')}
# RAM              : {info.get('ram_total','?')}
# Uptime           : {info.get('uptime','?')}

# ──────────────────────────────────────────────
# SECURITY STATUS
# ──────────────────────────────────────────────
# Firewall   : {fw_status}
# Antivirus  : {av_status}
# UAC        : {uac_s}

# ──────────────────────────────────────────────
# INCIDENT TYPE
# ──────────────────────────────────────────────
# [✔] Network Exposure
# [✔] Security Misconfiguration
# [ ] Malware Activity
# [ ] Unauthorized Access

# ──────────────────────────────────────────────
# OPEN PORTS & RISKS
# ──────────────────────────────────────────────
# {', '.join(str(p) for p in ports) if ports else 'No open ports detected'}

# ──────────────────────────────────────────────
# RISK ASSESSMENT
# ──────────────────────────────────────────────
# Overall Risk Level : {risk_label}

# ──────────────────────────────────────────────
# AI ANALYSIS
# ──────────────────────────────────────────────
# ML Model  : {'Active (Isolation Forest)' if AI_AVAILABLE and _model_trained else 'Training'}
# NLP Model : {'Active' if AI_AVAILABLE else 'Not Available'}

# ──────────────────────────────────────────────
# """

#     full_text.insert("end", report)
#     full_text.configure(state="disabled")

#     # ── Keep rest same ──
#     nlp_rep_text.configure(state="normal")
#     nlp_rep_text.delete("1.0","end")
#     nlp_rep_text.insert("end",
#         nlp_summarise_logs(state.get("events",[])) if state.get("events")
#         else "Run a scan to generate NLP log summary.")
#     nlp_rep_text.configure(state="disabled")

#     risk_val_lbl.configure(text=risk_label,fg=risk_col)


# ═══════════════════════════════════════════════════════════════════════════════
#  CORE SCAN
# ═══════════════════════════════════════════════════════════════════════════════
def do_scan(silent=False):
    if not silent:
        scan_btn.configure(state="disabled",text="  Scanning...  ")
        set_status("Collecting system information...")
    info=get_system_info(); state["info"]=info
    if not silent: populate_system_tab(info)
    append_syslog(f"[{ts()}] {'Auto' if silent else 'Manual'} scan started")
    write_log("Scan started")
    if not silent: set_status("Checking security configuration...")
    fw=check_firewall(); av=check_antivirus()
    uac=check_uac(); upd=check_windows_update()
    state["fw"]=fw; state["av"]=av
    if not silent: populate_security_tab(fw,av,uac,upd)
    append_raw_log(f"Firewall:{fw[0]}  AV:{av[0]}  UAC:{uac[0]}","ok" if fw[0]=="ENABLED" else "err")
    if not silent: set_status("Scanning ports...")
    ports=scan_ports(); state["ports"]=ports
    if not silent: populate_vuln_tab(ports)
    append_raw_log(f"Ports: {len(ports)} open","warn" if ports else "ok")
    if not silent: set_status("Analysing logs with NLP...")
    events=get_event_logs(); files=get_recent_files()
    state["events"]=events
    cat_counts={}
    if not silent:
        cat_counts=populate_log_tab(events,files)
        append_raw_log(f"NLP classified {len(events)} log entries","ok")
    else:
        for _,_,msg in events:
            cat,_,_=classify_log_entry(msg)
            cat_counts[cat]=cat_counts.get(cat,0)+1
    cpu_pct=float(info.get("cpu_usage",0))
    ram_pct=info.get("ram_pct",psutil.virtual_memory().percent)
    save_snapshot(cpu_pct,ram_pct,ports,fw[0],av[0])
    append_raw_log("Snapshot saved to database","ok")
    if AI_AVAILABLE:
        populate_ai_tab(cpu_pct,ram_pct,ports,fw[0],av[0],cat_counts)
    if not silent:
        set_status("Building report...")
        populate_report_tab(fw,av,ports,cat_counts)
        append_raw_log("Scan complete","ok")
        write_log("Scan complete")
        scan_btn.configure(state="normal",text="  Scan System  ")
        set_status(f"Scan complete — {datetime.datetime.now().strftime('%H:%M:%S')}")
    else:
        train_anomaly_model()
        append_syslog(f"[{ts()}] Auto-scan done — snapshot #{len(load_history())} saved")

def start_scan():
    threading.Thread(target=do_scan,daemon=True).start()

def schedule_auto_scan():
    threading.Thread(target=lambda: do_scan(silent=True),daemon=True).start()
    root.after(30*60*1000,schedule_auto_scan)
root.after(2*60*1000,schedule_auto_scan)

# #gemini new ////////////////////
# def perform_scan_logic():
#     # ... (After all port scanning and system checks are done) ...
    
#     # 1. Get the Risk Info and Raw Score
#     risk_info, raw_score = calculate_risk(state["fw"][0], state["av"][0], state["ports"])
#     status_text, color, bg = risk_info
    
#     # 2. Update UI
#     risk_val_lbl.config(text=status_text, fg=color)
#     risk_card.config(bg=bg)
#     risk_sub_lbl.config(text=f"Score: {raw_score} units", fg=color)
    
#     # 3. Trigger Forensic Snapshot if Risk is HIGH (Score >= 5)
#     # Note: Maine threshold 5 rakha hai kyunki aapka score 0-15 ke beech move karega
#     if raw_score >= 5:
#         log_msg = trigger_forensic_snapshot(raw_score) 
#         set_status(f"🚨 ALERT: {log_msg}")
#         messagebox.showwarning("Security Alert", f"High risk detected! Forensic evidence has been secured in 'Forensic_Snapshots' folder.")

# ═══════════════════════════════════════════════════════════════════════════════
#  REPORT EXPORT
# ═══════════════════════════════════════════════════════════════════════════════
def generate_report():
    switch_tab("Audit Report")
    if not state["info"]: start_scan(); return
    path=filedialog.asksaveasfilename(defaultextension=".docx",
                                       filetypes=[("Word Document","*.docx")],
                                       title="Save Audit Report")
    if not path: return
    info=state["info"]; fw=state["fw"]; av=state["av"]; ports=state["ports"]
    fw_status=fw[0]; av_status=av[0]
    risk_label,_,_=calculate_risk(fw_status,av_status,ports)
    ai_text=generate_ai_explanation(info,fw_status,av_status,ports,risk_label)
    uac_s,_=check_uac()

    doc=Document()
    doc.add_heading("Cyber Audit Report ",0)
    doc.add_paragraph(f"Generated: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_heading("System Information",1)
    for k,val in [("Hostname",info.get("hostname")),("IP",info.get("ip")),("OS",info.get("os")),
                   ("CPU",info.get("cpu_short")),("RAM",f"{info.get('ram_used')}/{info.get('ram_total')} GB"),
                   ("Uptime",info.get("uptime"))]:
        doc.add_paragraph(f"{k}: {val}")
    doc.add_heading("Disk Usage",1)
    for d in info.get("disks",[]):
        doc.add_paragraph(f"{d['device']} — {d['total']} GB ({d['percent']:.0f}% used)",style="List Bullet")
    doc.add_heading("Installed Software",1)
    for app in info.get("software",[]): doc.add_paragraph(app,style="List Bullet")
    doc.add_heading("Security Status",1)
    doc.add_paragraph(f"Firewall: {fw_status}\nAntivirus: {av_status}\nUAC: {uac_s}")
    doc.add_heading("Open Ports & Vulnerabilities",1)
    if ports:
        for p in ports:
            lvl,desc=PORT_RISK_INFO.get(p,("LOW",""))
            doc.add_paragraph(f"Port {p} ({COMMON_PORTS.get(p,'?')}) — {lvl}: {desc}",style="List Bullet")
    else:
        doc.add_paragraph("No open ports detected.")
    doc.add_heading("Risk Assessment",1)
    doc.add_paragraph(f"Overall Risk Level: {risk_label}")
    doc.add_heading("NLP Log Analysis",1)
    doc.add_paragraph(nlp_summarise_logs(state.get("events",[])))
    if AI_AVAILABLE and _model_trained:
        cpu_pct=float(info.get("cpu_usage",0))
        ram_pct=info.get("ram_pct",psutil.virtual_memory().percent)
        is_anom,score,expl=predict_anomaly(cpu_pct,ram_pct,ports,fw_status,av_status)
        doc.add_heading("AI Anomaly Detection (Isolation Forest)",1)
        doc.add_paragraph(f"Result: {'ANOMALY DETECTED' if is_anom else 'Normal'}")
        doc.add_paragraph(f"Score: {score:.4f}\n\n{expl}")
        rows=load_history(30)
        if rows:
            avg_cpu=sum(r[1] for r in rows)/len(rows)
            avg_ram=sum(r[2] for r in rows)/len(rows)
            doc.add_heading("Trend Summary (Last 30 Scans)",1)
            doc.add_paragraph(f"Avg CPU: {avg_cpu:.1f}%  |  Avg RAM: {avg_ram:.1f}%  |  Peak Risk: {max(r[7] for r in rows):.1f}")
    doc.add_heading("AI Threat Analysis",1)
    doc.add_paragraph(ai_text)
    doc.add_heading("Recommendations",1)
    hi_ports=[p for p in ports if p in HIGH_RISK_PORTS]
    if fw_status!="ENABLED": doc.add_paragraph("Enable Windows Firewall",style="List Bullet")
    if av_status!="ENABLED": doc.add_paragraph("Enable Antivirus Protection",style="List Bullet")
    for p in hi_ports: doc.add_paragraph(f"Close high-risk port {p} ({COMMON_PORTS.get(p,'')})",style="List Bullet")
    if not hi_ports and fw_status=="ENABLED" and av_status=="ENABLED":
        doc.add_paragraph("System appears secure. Maintain regular monitoring.",style="List Bullet")
    doc.save(path)
    messagebox.showinfo("Success","Report exported successfully!")

def refresh_ai():
    info=state["info"]; fw_s=state["fw"][0]; av_s=state["av"][0]; ports=state["ports"]
    cpu_pct=float(info.get("cpu_usage",0)) if info else psutil.cpu_percent()
    ram_pct=info.get("ram_pct",psutil.virtual_memory().percent) if info else psutil.virtual_memory().percent
    cat_counts={}
    for _,_,msg in state.get("events",[]):
        cat,_,_=classify_log_entry(msg); cat_counts[cat]=cat_counts.get(cat,0)+1
    populate_ai_tab(cpu_pct,ram_pct,ports,fw_s,av_s,cat_counts)

if AI_AVAILABLE:
    refresh_ai_btn.configure(command=refresh_ai)

scan_btn.configure(command=start_scan)
report_btn.configure(command=generate_report)
switch_tab("System Info")
root.mainloop()
